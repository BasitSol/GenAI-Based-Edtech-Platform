"""Teacher-facing DOCX/PDF exports for approved or reviewable assessments.

Exports are derived entirely from the persisted assessment content.  They never
call an LLM, re-run retrieval, or alter an assessment record, keeping a teacher
download reproducible and cost-free.
"""
from __future__ import annotations

from html import escape
from pathlib import Path
import re
from typing import Any, Literal

from backend.shared.core import PROCESSED_ROOT, slug


ExportFormat = Literal["docx", "pdf"]
_TABLE_SEPARATOR = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def export_directory() -> Path:
    """Keep transient teacher downloads outside source control and RAG builds."""
    path = PROCESSED_ROOT / "runtime" / "assessment_exports"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _filename(assessment: dict, extension: str, include_solutions: bool) -> Path:
    stem = f"assessment_{assessment['id']}_{slug(assessment['topic']) or 'untitled'}"
    audience = "teacher_key" if include_solutions else "student_paper"
    return export_directory() / f"{stem}_{audience}.{extension}"


def _questions(assessment: dict) -> list[dict]:
    questions = assessment.get("content", {}).get("questions", [])
    if not isinstance(questions, list) or not questions:
        raise ValueError("Assessment has no exportable questions.")
    return questions


def _table_cells(line: str) -> list[str]:
    """Split a GitHub-style Markdown table row into clean cell values."""
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    # Generated assessment tables contain simple curriculum text rather than
    # escaped pipes. Keeping the parser deliberately small makes export
    # behaviour predictable and independent of a browser Markdown renderer.
    return [cell.strip() for cell in value.split("|")]


def _implicit_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    """Recognise generated pipe tables that omit Markdown's separator row.

    LLMs commonly return truth tables as ``A B C | X`` followed by one binary
    row per line. That is clear tabular data to a human, but it is not valid
    Markdown and is consequently flattened by many UI and export renderers.
    This conservative repair requires at least a header plus one consecutive
    pipe-delimited row, so ordinary prose containing a pipe is left untouched.
    """
    candidates: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip() and "|" in lines[index]:
        candidates.append(lines[index])
        index += 1
    if len(candidates) < 2:
        return None

    rows = [_table_cells(line) for line in candidates]
    if not rows[0] or any(len(row) != len(rows[0]) for row in rows):
        return None

    # Expand compact truth tables such as ``A B C | X`` into four native
    # columns. Generic two-column tables (for example Field | Data type) keep
    # their explicit pipe-defined column boundaries.
    flattened = [re.findall(r"\S+", line.replace("|", " ")) for line in candidates]
    identifiers = re.compile(r"^[A-Za-z][A-Za-z0-9_]*$")
    binary = re.compile(r"^(?:0|1|true|false)$", re.IGNORECASE)
    if (2 <= len(flattened[0]) <= 10
            and all(identifiers.fullmatch(cell) for cell in flattened[0])
            and all(len(row) == len(flattened[0]) for row in flattened[1:])
            and all(binary.fullmatch(cell) for row in flattened[1:] for cell in row)):
        rows = flattened
    return rows, index


def structured_blocks(value: str) -> list[dict[str, Any]]:
    """Parse export-relevant Markdown into paragraphs, tables, and code.

    Web renderers can display these structures automatically, whereas DOCX and
    ReportLab do not. The persisted question remains unchanged; only the
    presentation layer converts its Markdown into native document elements.
    """
    lines = str(value or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        text = " ".join(part.strip() for part in paragraph if part.strip()).strip()
        if text:
            blocks.append({"type": "paragraph", "text": text})
        paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            blocks.append({"type": "code", "language": language, "text": "\n".join(code_lines).rstrip()})
            index += 1
            continue
        if (index + 1 < len(lines) and "|" in line
                and _TABLE_SEPARATOR.match(lines[index + 1])):
            flush_paragraph()
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines):
                candidate = lines[index]
                if "|" not in candidate or not candidate.strip():
                    break
                cells = _table_cells(candidate)
                if len(cells) != len(rows[0]):
                    break
                rows.append(cells)
                index += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        if "|" in line:
            implicit = _implicit_table(lines, index)
            if implicit:
                flush_paragraph()
                rows, index = implicit
                blocks.append({"type": "table", "rows": rows})
                continue
        if not stripped:
            flush_paragraph()
        elif (list_match := re.match(r"^\s*(?:([-*+])|(\d+)[.)])\s+(.+)$", line)):
            flush_paragraph()
            ordered = list_match.group(2) is not None
            items = [list_match.group(3).strip()]
            index += 1
            while index < len(lines):
                following = re.match(r"^\s*(?:([-*+])|(\d+)[.)])\s+(.+)$", lines[index])
                if not following or (following.group(2) is not None) != ordered:
                    break
                items.append(following.group(3).strip())
                index += 1
            blocks.append({"type": "list", "ordered": ordered, "items": items})
            continue
        else:
            paragraph.append(line)
        index += 1
    flush_paragraph()
    return blocks or [{"type": "paragraph", "text": str(value or "")}]


def structured_markdown(value: str) -> str:
    """Return browser-safe Markdown from the same canonical presentation AST.

    The dashboard and file exporters intentionally share ``structured_blocks``
    so a stored assessment has the same table/code semantics in every format.
    Existing assessments are repaired at display time and remain immutable.
    """
    rendered: list[str] = []
    for block in structured_blocks(value):
        if block["type"] == "paragraph":
            rendered.append(block["text"])
        elif block["type"] == "code":
            language = block.get("language", "")
            rendered.append(f"```{language}\n{block['text']}\n```")
        elif block["type"] == "table":
            rows = block["rows"]
            header = "| " + " | ".join(rows[0]) + " |"
            separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
            body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
            rendered.append("\n".join([header, separator, *body]))
        elif block["type"] == "list":
            marker = (lambda index: f"{index}.") if block["ordered"] else (lambda _index: "-")
            rendered.append("\n".join(
                f"{marker(index)} {item}" for index, item in enumerate(block["items"], 1)
            ))
    return "\n\n".join(rendered)


def _column_widths_dxa(rows: list[list[str]], total: int = 9360) -> list[int]:
    """Allocate fixed Word/ReportLab widths from bounded content demand."""
    if not rows or not rows[0]:
        return [total]
    columns = len(rows[0])
    demand = []
    for index in range(columns):
        longest = max(len(str(row[index])) for row in rows if index < len(row))
        demand.append(max(8, min(longest, 42)))
    minimum = 720
    available = total - minimum * columns
    scale = available / max(1, sum(demand))
    widths = [minimum + int(value * scale) for value in demand]
    widths[-1] += total - sum(widths)
    return widths


def _set_docx_run_font(run: Any, name: str = "Calibri", size: float | None = None,
                       color: str | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _apply_docx_numbering(paragraph: Any, num_id: int, level: int = 0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    num_properties = properties.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        properties.insert(0, num_properties)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_properties.extend([ilvl, num])


def _new_docx_numbering(document: Any, number_format: str, level_text: str,
                        *, left: int = 360, hanging: int = 360) -> int:
    """Create a real Word numbering definition using preset indent tokens."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), number_format)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left + hanging))
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left + hanging))
    indent.set(qn("w:hanging"), str(hanging))
    paragraph_properties.extend([tabs, indent])
    level.extend([start, fmt, text, suffix, paragraph_properties])
    abstract.append(multi)
    abstract.append(level)
    # OOXML requires every abstractNum before every num. Appending abstract
    # definitions after python-docx's built-in num entries looks plausible in
    # raw XML but Word ignores the out-of-order definitions and falls back to
    # bullets. Insert at the abstract/instance boundary instead.
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _set_docx_table_geometry(table: Any, rows: list[list[str]], widths: list[int]) -> None:
    """Set tblW/tblInd/tblGrid/tcW and cell margins to exact DXA values."""
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.insert(0, width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    for grid_column, column_width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(column_width))

    for row_index, row in enumerate(rows):
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(text) if text else " "
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.insert(0, cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[column_index]))
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = margins.find(qn(f"w:{edge}"))
                if margin is None:
                    margin = OxmlElement(f"w:{edge}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    _set_docx_run_font(run, size=10, bold=row_index == 0)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                cell_properties.append(shading)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            table.rows[0]._tr.get_or_add_trPr().append(repeat)


def _add_docx_blocks(document: Any, value: str, *, question_num_id: int | None = None) -> None:
    """Render paragraphs, code, and native fixed-geometry tables."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    blocks = structured_blocks(value)
    number_pending = question_num_id is not None
    for block in blocks:
        if block["type"] == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.1
            if number_pending:
                _apply_docx_numbering(paragraph, question_num_id)
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
                number_pending = False
            run = paragraph.add_run(block["text"])
            _set_docx_run_font(run, size=11)
        elif block["type"] == "table":
            if number_pending:
                paragraph = document.add_paragraph()
                _apply_docx_numbering(paragraph, question_num_id)
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
                number_pending = False
            rows = block["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            widths = _column_widths_dxa(rows)
            _set_docx_table_geometry(table, rows, widths)
            document.add_paragraph().paragraph_format.space_after = Pt(2)
        elif block["type"] == "code":
            paragraph = document.add_paragraph()
            if number_pending:
                _apply_docx_numbering(paragraph, question_num_id)
                paragraph.paragraph_format.keep_together = True
                number_pending = False
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.left_indent = Pt(8)
            paragraph.paragraph_format.right_indent = Pt(8)
            run = paragraph.add_run(block["text"])
            _set_docx_run_font(run, name="Consolas", size=9.5)
            properties = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F4F7")
            properties.append(shading)
        elif block["type"] == "list":
            if number_pending:
                paragraph = document.add_paragraph()
                _apply_docx_numbering(paragraph, question_num_id)
                paragraph.paragraph_format.keep_with_next = True
                number_pending = False
            num_id = _new_docx_numbering(
                document,
                "decimal" if block["ordered"] else "bullet",
                "%1." if block["ordered"] else "\u2022",
            )
            for item in block["items"]:
                paragraph = document.add_paragraph()
                _apply_docx_numbering(paragraph, num_id)
                paragraph.paragraph_format.space_after = Pt(3)
                if question_num_id is not None:
                    paragraph.paragraph_format.keep_with_next = True
                run = paragraph.add_run(item)
                _set_docx_run_font(run, size=11)
    if number_pending:
        paragraph = document.add_paragraph()
        _apply_docx_numbering(paragraph, question_num_id)


def _pdf_blocks(value: str, styles: Any, *, number: int | None = None) -> list[Any]:
    """Render structured content as ReportLab flowables."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import ListFlowable, ListItem, LongTable, Paragraph, Spacer, Table, TableStyle, XPreformatted

    flowables: list[Any] = []
    prefix_pending = number is not None
    for block in structured_blocks(value):
        if block["type"] == "paragraph":
            prefix = f"<b>{number}.</b> " if prefix_pending else ""
            prefix_pending = False
            flowables.append(Paragraph(prefix + escape(block["text"]), styles["BodyText"]))
        elif block["type"] == "table":
            if prefix_pending:
                flowables.append(Paragraph(f"<b>{number}.</b>", styles["BodyText"]))
                prefix_pending = False
            rows = [[Paragraph(escape(cell) if cell else "&#160;", styles["BodyText"]) for cell in row]
                    for row in block["rows"]]
            widths_dxa = _column_widths_dxa(block["rows"])
            column_widths = [16.5 * cm * value / sum(widths_dxa) for value in widths_dxa]
            table = LongTable(rows, colWidths=column_widths, repeatRows=1,
                              minRowHeights=[0.65 * cm] * len(rows), splitByRow=1)
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#667085")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            flowables.extend([Spacer(1, 0.12 * cm), table])
        elif block["type"] == "code":
            prefix = f"<b>{number}.</b><br/>" if prefix_pending else ""
            prefix_pending = False
            code_flowables: list[Any] = []
            if prefix:
                code_flowables.append(Paragraph(prefix, styles["BodyText"]))
            code_flowables.append(XPreformatted(block["text"], styles["Code"]))
            box = Table([[code_flowables]], colWidths=[16.5 * cm])
            box.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F2F4F7")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#98A2B3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]))
            flowables.append(box)
        elif block["type"] == "list":
            if prefix_pending:
                flowables.append(Paragraph(f"<b>{number}.</b>", styles["BodyText"]))
                prefix_pending = False
            items = [ListItem(Paragraph(escape(item), styles["BodyText"])) for item in block["items"]]
            list_options = {
                "bulletType": "1" if block["ordered"] else "bullet",
                "leftIndent": 18,
                "bulletFontName": "Helvetica",
                "bulletFontSize": 9,
            }
            if block["ordered"]:
                list_options.update({"start": "1", "bulletFormat": "%s."})
            flowables.append(ListFlowable(items, **list_options))
    if prefix_pending:
        flowables.append(Paragraph(f"<b>{number}.</b>", styles["BodyText"]))
    return flowables


def _docx_field(paragraph: Any, instruction: str) -> None:
    """Insert a native Word field so page numbering updates in Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field = OxmlElement("w:instrText")
    field.set(qn("xml:space"), "preserve")
    field.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run.extend([begin, field, separate, placeholder, end])
    paragraph._p.append(run)


def _configure_docx_styles(document: Any) -> None:
    """Apply the standard-business-brief typography preset."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Title", 22, "17365D", 0, 8),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 10, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Assessment Metadata" not in styles:
        metadata = styles.add_style("Assessment Metadata", WD_STYLE_TYPE.PARAGRAPH)
        metadata.font.name = "Calibri"
        metadata.font.size = Pt(9.5)
        metadata.font.color.rgb = RGBColor.from_string("475467")
        metadata.paragraph_format.space_after = Pt(3)


def _docx_header_footer(document: Any, title: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    section = document.sections[0]
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"COMPUTER SCIENCE  |  {title}")
    _set_docx_run_font(run, size=8.5, color="667085", bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Computer Science Assessment  •  Page ")
    _set_docx_run_font(run, size=8.5, color="667085")
    _docx_field(footer, " PAGE ")
    run = footer.add_run(" of ")
    _set_docx_run_font(run, size=8.5, color="667085")
    _docx_field(footer, " NUMPAGES ")


def _docx_callout(document: Any, heading: str, text: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = document.add_table(rows=1, cols=1)
    _set_docx_table_geometry(table, [[" "]], [9360])
    cell = table.cell(0, 0)
    cell.text = ""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EAF2F8")
    properties.append(shading)
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    properties.append(borders)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(heading.upper())
    _set_docx_run_font(run, size=9, color="1F4D78", bold=True)
    paragraph = cell.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        _set_docx_run_font(run, size=10.5)


def _docx_answer_lines(document: Any, count: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for _ in range(max(1, min(count, 5))):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(8)
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "3")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D0D5DD")
        borders.append(bottom)
        paragraph._p.get_or_add_pPr().append(borders)


def _export_docx(assessment: dict, include_solutions: bool) -> Path:
    """Create an editable, print-ready examination paper and teacher key."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    _configure_docx_styles(document)

    content = assessment.get("content", {})
    title_text = content.get("title") or assessment["topic"]
    _docx_header_footer(document, title_text)
    document.core_properties.title = title_text
    document.core_properties.subject = "Computer Science educational assessment"
    document.core_properties.author = "Computer Science RAG"

    eyebrow = document.add_paragraph(style="Assessment Metadata")
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("COMPUTER SCIENCE ASSESSMENT")
    _set_docx_run_font(run, size=9, color="2E74B5", bold=True)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)
    metadata = document.add_paragraph(style="Assessment Metadata")
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status = "APPROVED" if assessment.get("status") == "approved" else "TEACHER-REVIEW DRAFT"
    metadata.add_run(
        f"{assessment['type'].replace('_', ' ').title()}  •  "
        f"{assessment['difficulty'].title()}  •  {status}"
    )

    identity = document.add_table(rows=1, cols=3)
    _set_docx_table_geometry(
        identity,
        [["Candidate name: ____________________", "Class: ______________", "Date: ______________"]],
        [4680, 2340, 2340],
    )
    for cell in identity.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F8FAFC")
        cell._tc.get_or_add_tcPr().append(shading)

    instructions = content.get(
        "instructions",
        "Answer all questions. Show working where appropriate.",
    )
    document.add_heading("Instructions", level=1)
    _docx_callout(document, "Before you begin", instructions)
    document.add_heading("Questions", level=1)
    question_num_id = _new_docx_numbering(document, "decimal", "%1.", left=360, hanging=360)
    for index, question in enumerate(_questions(assessment), 1):
        _add_docx_blocks(document, question["question"], question_num_id=question_num_id)
        if question.get("question_type") == "MCQ":
            option_num_id = _new_docx_numbering(document, "upperLetter", "%1.", left=720, hanging=360)
            for option in question.get("options", []):
                paragraph = document.add_paragraph()
                _apply_docx_numbering(paragraph, option_num_id)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.keep_with_next = True
                run = paragraph.add_run(str(option))
                _set_docx_run_font(run, size=11)
        if document.paragraphs:
            document.paragraphs[-1].paragraph_format.keep_with_next = True
        marks = int(question["marks"])
        mark_line = document.add_paragraph()
        mark_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mark_line.paragraph_format.space_before = Pt(3)
        mark_line.paragraph_format.space_after = Pt(3)
        mark_line.paragraph_format.keep_with_next = True
        run = mark_line.add_run(f"[{marks} {'mark' if marks == 1 else 'marks'}]")
        _set_docx_run_font(run, size=9.5, color="1F4D78", bold=True)
        _docx_answer_lines(document, 2 if question.get("question_type") == "MCQ" else marks)

    if include_solutions:
        document.add_page_break()
        document.add_heading("Teacher solution key and rubric", level=1)
        _docx_callout(
            document,
            "Teacher use",
            "Award marks using the criteria below. Accept equivalent technically correct wording.",
        )
        for index, question in enumerate(_questions(assessment), 1):
            number = question.get("number", index)
            document.add_heading(f"Question {number}  |  {question['marks']} marks", level=2)
            if question.get("question_type") == "MCQ":
                correct = question.get("correct_option", "")
                options = question.get("options", [])
                label = chr(65 + options.index(correct)) if correct in options else ""
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"Correct answer: {label}. {correct}")
                _set_docx_run_font(run, size=11, color="1F4D78", bold=True)
                paragraph.paragraph_format.keep_with_next = True
            answer_paragraph_start = len(document.paragraphs)
            _add_docx_blocks(document, question["model_answer"])
            has_table = any(
                block["type"] == "table" for block in structured_blocks(question["model_answer"])
            )
            if not has_table:
                for paragraph in document.paragraphs[answer_paragraph_start:]:
                    paragraph.paragraph_format.keep_with_next = True
            heading = document.add_paragraph()
            heading.paragraph_format.keep_with_next = True
            run = heading.add_run("Marking points")
            _set_docx_run_font(run, size=10, color="475467", bold=True)
            rubric_num_id = _new_docx_numbering(document, "bullet", "\u2022")
            rubric_points = question.get("rubric", [])
            for point_index, point in enumerate(rubric_points):
                paragraph = document.add_paragraph()
                _apply_docx_numbering(paragraph, rubric_num_id)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.keep_with_next = point_index < len(rubric_points) - 1
                run = paragraph.add_run(str(point))
                _set_docx_run_font(run, size=10.5)

    # Ask Word to refresh PAGE/NUMPAGES fields when the document opens.
    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    path = _filename(assessment, "docx", include_solutions)
    document.save(path)
    return path


def _export_pdf(assessment: dict, include_solutions: bool) -> Path:
    """Create a polished ReportLab examination paper with a teacher key."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable,
        CondPageBreak,
        KeepTogether,
        ListFlowable,
        ListItem,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    path = _filename(assessment, "pdf", include_solutions)
    content = assessment.get("content", {})
    title_text = content.get("title") or assessment["topic"]
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Helvetica"
    styles["Normal"].fontSize = 10.5
    styles["Normal"].leading = 14
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 10.5
    styles["BodyText"].leading = 14
    styles["BodyText"].spaceAfter = 5
    styles["Title"].fontName = "Helvetica-Bold"
    styles["Title"].fontSize = 22
    styles["Title"].leading = 26
    styles["Title"].textColor = colors.HexColor("#17365D")
    styles["Title"].alignment = TA_CENTER
    styles["Title"].spaceAfter = 8
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading1"].fontSize = 16
    styles["Heading1"].leading = 20
    styles["Heading1"].textColor = colors.HexColor("#2E74B5")
    styles["Heading1"].spaceBefore = 14
    styles["Heading1"].spaceAfter = 8
    styles["Heading2"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontSize = 12.5
    styles["Heading2"].leading = 16
    styles["Heading2"].textColor = colors.HexColor("#1F4D78")
    styles["Heading2"].spaceBefore = 10
    styles["Heading2"].spaceAfter = 5
    styles["Code"].fontName = "Courier"
    styles["Code"].fontSize = 8.5
    styles["Code"].leading = 11
    styles.add(ParagraphStyle(
        name="Eyebrow",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#2E74B5"),
        alignment=TA_CENTER,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="Metadata",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#475467"),
        alignment=TA_CENTER,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="Mark",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        textColor=colors.HexColor("#1F4D78"),
        alignment=TA_RIGHT,
    ))

    document = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.88 * inch,
        bottomMargin=0.72 * inch,
        title=title_text,
        author="Computer Science RAG",
        subject="Computer Science educational assessment",
    )

    def page_chrome(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D0D5DD"))
        canvas.setLineWidth(0.5)
        canvas.line(0.75 * inch, letter[1] - 0.56 * inch, letter[0] - 0.75 * inch, letter[1] - 0.56 * inch)
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.HexColor("#667085"))
        canvas.drawString(0.75 * inch, letter[1] - 0.43 * inch, "COMPUTER SCIENCE")
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(letter[0] - 0.75 * inch, letter[1] - 0.43 * inch, title_text[:70])
        canvas.line(0.75 * inch, 0.48 * inch, letter[0] - 0.75 * inch, 0.48 * inch)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(0.75 * inch, 0.31 * inch, "Computer Science Assessment")
        canvas.drawRightString(letter[0] - 0.75 * inch, 0.31 * inch, f"Page {doc.page}")
        canvas.restoreState()

    status = "APPROVED" if assessment.get("status") == "approved" else "TEACHER-REVIEW DRAFT"
    identity = Table(
        [[
            Paragraph("<b>Candidate name:</b> ____________________", styles["BodyText"]),
            Paragraph("<b>Class:</b> ______________", styles["BodyText"]),
            Paragraph("<b>Date:</b> ______________", styles["BodyText"]),
        ]],
        colWidths=[3.25 * inch, 1.65 * inch, 1.6 * inch],
    )
    identity.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#EAECF0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    instructions = content.get(
        "instructions",
        "Answer all questions. Show working where appropriate.",
    )
    callout = Table(
        [[Paragraph("<b>BEFORE YOU BEGIN</b><br/>" + escape(instructions), styles["BodyText"])]],
        colWidths=[6.5 * inch],
    )
    callout.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF2F8")),
        ("LINEBEFORE", (0, 0), (0, -1), 4, colors.HexColor("#2E74B5")),
        ("LEFTPADDING", (0, 0), (-1, -1), 11),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story: list[Any] = [
        Paragraph("COMPUTER SCIENCE ASSESSMENT", styles["Eyebrow"]),
        Paragraph(escape(title_text), styles["Title"]),
        Paragraph(
            escape(
                f"{assessment['type'].replace('_', ' ').title()}  •  "
                f"{assessment['difficulty'].title()}  •  {status}"
            ),
            styles["Metadata"],
        ),
        identity,
        Paragraph("Instructions", styles["Heading1"]),
        callout,
        Paragraph("Questions", styles["Heading1"]),
    ]
    for index, question in enumerate(_questions(assessment), 1):
        number = question.get("number", index)
        story.extend(_pdf_blocks(question["question"], styles, number=number))
        if question.get("question_type") == "MCQ":
            options = [
                ListItem(Paragraph(escape(str(option)), styles["BodyText"]))
                for option in question.get("options", [])
            ]
            story.append(ListFlowable(
                options,
                bulletType="A",
                start="A",
                bulletFormat="%s.",
                leftIndent=32,
                bulletFontName="Helvetica-Bold",
                bulletFontSize=9,
            ))
        marks = int(question["marks"])
        story.append(Paragraph(f"[{marks} {'mark' if marks == 1 else 'marks'}]", styles["Mark"]))
        for _ in range(2 if question.get("question_type") == "MCQ" else max(1, min(marks, 5))):
            story.extend([
                Spacer(1, 9),
                HRFlowable(width="100%", thickness=0.35, color=colors.HexColor("#D0D5DD")),
            ])
        story.append(Spacer(1, 9))

    if include_solutions:
        story.extend([
            PageBreak(),
            Paragraph("Teacher solution key and rubric", styles["Heading1"]),
            Table(
                [[Paragraph(
                    "<b>TEACHER USE</b><br/>Award marks using the criteria below. "
                    "Accept equivalent technically correct wording.",
                    styles["BodyText"],
                )]],
                colWidths=[6.5 * inch],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF7E6")),
                    ("LINEBEFORE", (0, 0), (0, -1), 4, colors.HexColor("#D97706")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 11),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]),
            ),
        ])
        for index, question in enumerate(_questions(assessment), 1):
            number = question.get("number", index)
            solution: list[Any] = [Paragraph(
                escape(f"Question {number}  |  {question['marks']} marks"),
                styles["Heading2"],
            )]
            if question.get("question_type") == "MCQ":
                correct = question.get("correct_option", "")
                options = question.get("options", [])
                label = chr(65 + options.index(correct)) if correct in options else ""
                solution.append(Paragraph(
                    f"<b>Correct answer:</b> {escape(label)}. {escape(correct)}",
                    styles["BodyText"],
                ))
            solution.extend(_pdf_blocks(question["model_answer"], styles))
            solution.append(Paragraph("<b>Marking points</b>", styles["BodyText"]))
            rubric = [
                ListItem(Paragraph(escape(str(point)), styles["BodyText"]))
                for point in question.get("rubric", [])
            ]
            if rubric:
                solution.append(ListFlowable(
                    rubric,
                    bulletType="bullet",
                    leftIndent=20,
                    bulletFontName="Helvetica",
                    bulletFontSize=9,
                ))
            story.append(CondPageBreak(1.35 * inch))
            model_blocks = structured_blocks(question["model_answer"])
            if any(block["type"] == "table" for block in model_blocks):
                story.extend(solution)
            else:
                story.append(KeepTogether(solution))
    document.build(story, onFirstPage=page_chrome, onLaterPages=page_chrome)
    return path


def export_assessment(assessment: dict, output_format: ExportFormat, include_solutions: bool = True) -> Path:
    """Export persisted content to Word or PDF with no hidden generation step."""
    if output_format == "docx":
        return _export_docx(assessment, include_solutions)
    if output_format == "pdf":
        return _export_pdf(assessment, include_solutions)
    raise ValueError(f"Unsupported assessment export format: {output_format}")
